"""
Document Generation Engine for SGA Workcover Dashboard.
Generates pre-filled .docx files for workcover case management.
"""

import io
from datetime import datetime, date, timedelta
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT


# ── Constants ──────────────────────────────────────────────────

RTW_COORDINATOR = {
    "name": "Ben North",
    "role": "Return to Work Coordinator",
    "phone": "0403 427 790",
    "email": "Ben.n@sga.com.au",
    "address": "8 Guest Street, Hawthorn 3122, VIC",
}

EMPLOYER = {
    "name": "Sanikleen Group Australia Pty Ltd",
    "short": "Sanikleen Group Australia",
}

# Insurer / Agent by state
AGENTS = {
    "VIC": {"name": "DXC Technology", "phone": "1300 365 885", "web": "www.dxc.com", "address": "GPO Box 4028, Melbourne VIC 3001"},
    "NSW": {"name": "Allianz", "phone": "13 10 13", "web": "www.allianz.com.au", "address": "GPO Box 4049, Sydney NSW 2001"},
    "QLD": {"name": "WorkCover Queensland", "phone": "1300 362 128", "web": "www.workcoverqld.com.au", "address": "GPO Box 2459, Brisbane QLD 4001"},
}

# Suitable duties levels
SUITABLE_DUTIES = {
    1: {
        "title": "Level 1 - Seated / Observational Duties (Initial Capacity)",
        "purpose": "Maintain workplace engagement while minimising physical demand and preventing aggravation of injury",
        "duties": [
            "Seated training (theoretical, procedural, or safety-based)",
            "On-site walkthroughs for observation purposes only",
            "Job observation and task familiarisation",
            "Supervision and oversight of other employees",
            "Quality assurance checks and error identification",
            "Review of procedures, inductions, and Safe Work Method Statements",
            "Administrative or reporting tasks related to cleaning operations",
        ],
        "restrictions": [
            "No cleaning duties",
            "No lifting, pushing, pulling, or carrying",
            "No repetitive movements",
            "No prolonged standing or walking",
            "No use of tools, machinery, or chemicals",
        ],
        "hours_min": 3,
        "rest_break": "15 minutes at regular 1-hour intervals",
    },
    2: {
        "title": "Level 2 - Modified Duties (Low Physical Demand)",
        "purpose": "Introduce very light, controlled physical activity while remaining within medical restrictions",
        "duties": [
            "Wiping down and drying machinery and equipment",
            "Cleaning and drying benchtops and surfaces",
            "Light scrubbing and scouring tasks",
            "Waterproofing machinery",
            "Amenities cleaning - replacing bin liners, cleaning toilets/sinks/amenities fixtures",
            "General cleaning tasks that can be performed one-handed, at waist height, or with minimal standing",
        ],
        "restrictions": [
            "No heavy scrubbing",
            "No lifting or carrying of items",
            "No bending below knee height",
            "No overhead work",
            "No dismantling of machinery",
            "No mopping",
            "No use of high-pressure equipment",
        ],
        "hours_min": 3,
        "rest_break": "15 minutes at regular 1-hour intervals",
    },
    3: {
        "title": "Level 3 - Modified Duties (Moderate Physical Demand)",
        "purpose": "Progressively rebuild functional capacity and tolerance to work activities",
        "duties": [
            "All duties listed under Level 2",
            "Dismantling machinery prior to cleaning (e.g., removal of plate covers)",
            "Carrying chemicals or equipment up to approximately 0-7 kilograms",
            "Heavier scrubbing tasks",
            "Cleaning of lower sections of machinery and equipment",
            "Mopping / sweeping of floors",
            "Increased range of movement and coverage of cleaning areas",
        ],
        "restrictions": [
            "No use of high-pressure hoses",
            "No confined space entry",
            "No working at heights",
            "No prolonged repetitive tasks without appropriate rest breaks",
            "No lifting beyond medically certified limits",
        ],
        "hours_min": 3,
        "rest_break": "15 minutes at regular 2-hour intervals",
    },
    4: {
        "title": "Level 4 - Pre-Injury Duties (Full Capacity)",
        "purpose": "Return to full duties consistent with worker's pre-injury role",
        "duties": [
            "All standard cleaning duties as per employee's position description",
            "Use of high-pressure hoses",
            "High cleaning tasks (where required)",
            "Confined space cleaning (where required)",
            "Full manual handling tasks",
            "Full use of tools, equipment, and chemicals",
            "Normal shift duration, workload, and pace",
        ],
        "restrictions": ["Nil, other than standard workplace health and safety requirements"],
        "hours_min": None,
        "rest_break": None,
    },
}


# ── Helpers ────────────────────────────────────────────────────

def _fmt(val, fallback="[REQUIRED]"):
    """Return value if truthy, else a marker."""
    if val is None or (isinstance(val, str) and val.strip() == ""):
        return fallback
    return str(val)


def _fmt_date(val, fallback="[REQUIRED]"):
    """Format a date string or date object."""
    if not val:
        return fallback
    if isinstance(val, (date, datetime)):
        return val.strftime("%d/%m/%Y")
    try:
        return datetime.strptime(str(val), "%Y-%m-%d").strftime("%d/%m/%Y")
    except (ValueError, TypeError):
        return str(val)


def _set_cell_text(cell, text, bold=False, size=10, alignment=None):
    """Set text in a table cell."""
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.bold = bold
    if alignment:
        p.alignment = alignment


def _add_heading_run(paragraph, text, bold=True, size=12, color=None):
    """Add a formatted run to a paragraph."""
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run


def _marker_run(paragraph, text):
    """Add a yellow-highlighted marker run."""
    run = paragraph.add_run(text)
    run.bold = True
    run.font.color.rgb = RGBColor(0xCC, 0x66, 0x00)
    return run


def _add_paragraph(doc, text="", style=None, bold=False, size=11, space_after=6, alignment=None):
    """Add a paragraph with formatting."""
    p = doc.add_paragraph(style=style)
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(space_after)
    if alignment:
        p.alignment = alignment
    return p


def _set_table_style(table):
    """Apply consistent styling to a table."""
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)


def _get_suitable_level(capacity):
    """Determine suitable duties level based on capacity."""
    if not capacity:
        return 1
    cap = capacity.lower()
    if "no capacity" in cap:
        return 1
    elif "modified" in cap:
        return 2
    elif "full" in cap or "clearance" in cap:
        return 4
    return 2


def _build_progressive_hours(current_hours, pre_injury_hours, weeks=4):
    """Build a 4-week progressive schedule."""
    if not current_hours or current_hours <= 0:
        current_hours = 3
    if not pre_injury_hours or pre_injury_hours <= 0:
        pre_injury_hours = 38

    step = (pre_injury_hours - current_hours) / max(weeks - 1, 1)
    schedule = []
    for w in range(weeks):
        hrs = min(current_hours + (step * w), pre_injury_hours)
        schedule.append(round(hrs, 1))
    return schedule


# ── Document Generators ───────────────────────────────────────

def generate_letter_to_worker(case_data):
    """
    Generate Letter to Worker regarding Return to Work Arrangements.
    Returns a BytesIO object containing the .docx file.
    """
    doc = Document()

    # Set narrow margins
    for section in doc.sections:
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Header info block
    table = doc.add_table(rows=4, cols=2)
    _set_cell_text(table.cell(0, 0), "WORKER NAME:", bold=True)
    _set_cell_text(table.cell(0, 1), _fmt(case_data.get("worker_name")))
    _set_cell_text(table.cell(1, 0), "CLAIM NUMBER:", bold=True)
    _set_cell_text(table.cell(1, 1), _fmt(case_data.get("claim_number")))
    _set_cell_text(table.cell(2, 0), "DATE:", bold=True)
    _set_cell_text(table.cell(2, 1), date.today().strftime("%d/%m/%Y"))
    _set_cell_text(table.cell(3, 0), "RTW COORDINATOR:", bold=True)
    _set_cell_text(table.cell(3, 1), RTW_COORDINATOR["name"])
    _set_table_style(table)

    doc.add_paragraph("")

    # Title
    _add_paragraph(doc, "Recovery and Return to Work", bold=True, size=14,
                   alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # Body
    worker_first = (case_data.get("worker_name") or "").split()[0] if case_data.get("worker_name") else "[REQUIRED]"

    _add_paragraph(doc,
        f"Dear {worker_first},",
        size=11, space_after=12)

    _add_paragraph(doc,
        f"{EMPLOYER['short']} is committed to supporting your recovery and return to safe, "
        f"suitable and sustainable employment following your workplace injury. "
        f"We have developed Return to Work Arrangements based on available medical information "
        f"and in consultation with you and your treating health practitioner.",
        size=11, space_after=12)

    _add_paragraph(doc, "What do I need to do?", bold=True, size=12, space_after=6)

    actions = [
        "Please read the Return to Work Arrangements attached",
        "Discuss the Return to Work Arrangements with your doctor",
        "Provide feedback on the Return to Work Arrangements",
        "Sign second page if satisfied and return a copy",
    ]
    for action in actions:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(action)
        run.font.size = Pt(11)

    doc.add_paragraph("")

    _add_paragraph(doc,
        "A copy of this letter has been sent to your doctor(s).",
        size=11, space_after=12)

    _add_paragraph(doc,
        "If you have any questions or concerns, please don't hesitate to contact me.",
        size=11, space_after=12)

    _add_paragraph(doc, "Kind regards,", size=11, space_after=6)

    doc.add_paragraph("")
    _add_paragraph(doc, RTW_COORDINATOR["name"], bold=True, size=11, space_after=2)
    _add_paragraph(doc, RTW_COORDINATOR["role"], size=10, space_after=2)
    _add_paragraph(doc, f"Phone: {RTW_COORDINATOR['phone']}", size=10, space_after=2)
    _add_paragraph(doc, f"Email: {RTW_COORDINATOR['email']}", size=10, space_after=2)
    _add_paragraph(doc, RTW_COORDINATOR["address"], size=10, space_after=2)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def generate_rtw_information(case_data):
    """
    Generate Important Return to Work Information document (VIC template).
    Returns BytesIO with .docx.
    """
    doc = Document()
    state = case_data.get("state", "VIC")
    agent = AGENTS.get(state, AGENTS["VIC"])
    employer_name = EMPLOYER["short"]

    for section in doc.sections:
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # Title
    _add_paragraph(doc, "Important Return to Work Information", bold=True, size=16,
                   alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    # Section 1 - Employer obligations
    _add_paragraph(doc,
        f"{employer_name} return to work obligations under Workers' Compensation legislation",
        bold=True, size=12, space_after=6)

    _add_paragraph(doc,
        f"{employer_name} has obligations under workers' compensation legislation. "
        f"This document outlines how {employer_name} will meet these obligations and your rights "
        f"and obligations as a worker.",
        size=10, space_after=10)

    # Section 2 - How employer meets obligations
    _add_paragraph(doc,
        f"How {employer_name} will meet its obligations",
        bold=True, size=12, space_after=6)

    obligations = [
        ("Make return to work information available",
         f"{employer_name} will make return to work information available to workers about:\n"
         f"- The obligations of {employer_name} under the legislation\n"
         f"- The rights and obligations of workers under the legislation\n"
         f"- The name and contact details of the authorised Agent\n"
         f"- The name and contact details of the Return to Work Coordinator\n"
         f"- The procedure for resolving return to work issues"),
        ("Provide employment",
         f"{employer_name} will provide suitable employment to an injured worker (if they have current work capacity) "
         f"or pre-injury employment (if no longer incapacitated) for 52 weeks of incapacity from the date of the "
         f"Certificate of Capacity or Worker's Injury Claim Form."),
        ("Plan return to work",
         f"{employer_name} will commence return to work planning from receipt of the Worker's Injury Claim Form "
         f"or initial Certificate of Capacity. As part of planning, {employer_name} will:\n"
         f"- Obtain relevant information about the injured worker's capacity for work\n"
         f"- Consider reasonable workplace support, aids or modifications\n"
         f"- Assess and propose options for suitable or pre-injury employment\n"
         f"- Engage in consultation about return to work\n"
         f"- Provide clear, accurate and current details of return to work arrangements\n"
         f"- Monitor the worker's progress as often as necessary"),
        ("Consult about the return to work of a worker",
         f"{employer_name} will consult with the worker, treating health practitioner (with consent), "
         f"and occupational rehabilitation provider (if involved). "
         f"The worker may be represented, assisted and supported during the return to work process."),
        ("Nominate and appoint a Return to Work Coordinator",
         f"{employer_name} has nominated and appointed a Return to Work Coordinator at all times, "
         f"who has the appropriate level of seniority and is competent to assist the employer meet its obligations."),
    ]

    for heading, body in obligations:
        _add_paragraph(doc, heading, bold=True, size=11, space_after=4)
        _add_paragraph(doc, body, size=10, space_after=10)

    # Section 3 - Worker rights and obligations
    _add_paragraph(doc, "Worker's return to work rights and obligations", bold=True, size=12, space_after=6)

    _add_paragraph(doc, "Injured worker rights:", bold=True, size=11, space_after=4)
    rights = [
        "Be provided with return to work information and be consulted about how that information is made available",
        "Be provided with suitable employment (if you have current work capacity) or pre-injury employment (if no longer incapacitated) for 52 weeks",
        "Be consulted about planning return to work",
        "Be provided with clear, accurate and current details of return to work arrangements",
        "Be represented, assisted and supported during any stage of the return to work process",
    ]
    for r in rights:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(r)
        run.font.size = Pt(10)

    _add_paragraph(doc, "", space_after=6)
    _add_paragraph(doc, "Injured worker's obligations:", bold=True, size=11, space_after=4)
    obligations_list = [
        "Make reasonable efforts to actively participate and cooperate in planning for return to work",
        "Make reasonable efforts to return to work in suitable or pre-injury employment",
        "Actively use occupational rehabilitation services if provided",
        "Actively participate and cooperate in assessments of capacity and rehabilitation progress",
        "Attempt to resolve return to work issues in accordance with the agreed procedure",
    ]
    for o in obligations_list:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(o)
        run.font.size = Pt(10)

    _add_paragraph(doc, "", space_after=6)
    _add_paragraph(doc,
        "If a worker does not comply with these obligations, weekly payments may be suspended, "
        "terminated or ceased.",
        size=10, space_after=12)

    # Section 4 - Contact details
    _add_paragraph(doc, "Where to get help", bold=True, size=12, space_after=6)

    # RTW Coordinator contact
    _add_paragraph(doc, "Our Return to Work Coordinator:", bold=True, size=11, space_after=4)
    _add_paragraph(doc, f"Name: {RTW_COORDINATOR['name']}", size=10, space_after=2)
    _add_paragraph(doc, f"Phone: {RTW_COORDINATOR['phone']}", size=10, space_after=2)
    _add_paragraph(doc, f"Email: {RTW_COORDINATOR['email']}", size=10, space_after=2)
    _add_paragraph(doc, f"Address: {RTW_COORDINATOR['address']}", size=10, space_after=10)

    # Agent contact
    _add_paragraph(doc, "Our Authorised Agent:", bold=True, size=11, space_after=4)
    _add_paragraph(doc, f"Name: {agent['name']}", size=10, space_after=2)
    _add_paragraph(doc, f"Phone: {agent['phone']}", size=10, space_after=2)
    _add_paragraph(doc, f"Web: {agent['web']}", size=10, space_after=2)
    _add_paragraph(doc, f"Address: {agent['address']}", size=10, space_after=10)

    # WorkSafe contact
    if state == "VIC":
        _add_paragraph(doc, "WorkSafe Victoria:", bold=True, size=11, space_after=4)
        _add_paragraph(doc, "Phone: 1800 136 089", size=10, space_after=2)
        _add_paragraph(doc, "Web: worksafe.vic.gov.au", size=10, space_after=2)
        _add_paragraph(doc, "Email: info@worksafe.vic.gov.au", size=10, space_after=2)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def generate_register_of_injury(case_data, incident_data=None):
    """
    Generate Register of Injury and Investigation form.
    Returns BytesIO with .docx.
    """
    doc = Document()
    incident = incident_data or {}

    for section in doc.sections:
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    _add_paragraph(doc, "REGISTER OF INJURY AND INVESTIGATION", bold=True, size=14,
                   alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    # ── PART A - INCIDENT DETAILS ──
    _add_paragraph(doc, "PART A - INCIDENT DETAILS", bold=True, size=12, space_after=6)

    table_a = doc.add_table(rows=10, cols=2)
    table_a.style = "Table Grid"
    fields_a = [
        ("Employee Name", _fmt(case_data.get("worker_name"))),
        ("Workplace / Site", _fmt(case_data.get("site"))),
        ("Date of Birth", _fmt(incident.get("dob"), "[REQUIRED]")),
        ("Occupation", _fmt(incident.get("occupation"), "[REVIEW] Cleaner")),
        ("Date of Incident", _fmt_date(case_data.get("date_of_injury"))),
        ("Date Reported", _fmt(incident.get("date_reported"), "[REQUIRED]")),
        ("Task being performed", _fmt(incident.get("task_performed"), "[REQUIRED]")),
        ("Location where accident occurred", _fmt(incident.get("location_detail"), "[REQUIRED]")),
        ("What happened?", _fmt(case_data.get("injury_description"))),
        ("Witnesses", _fmt(incident.get("witnesses"), "[REQUIRED]")),
    ]
    for i, (label, val) in enumerate(fields_a):
        _set_cell_text(table_a.cell(i, 0), label, bold=True, size=10)
        _set_cell_text(table_a.cell(i, 1), val, size=10)
    _set_table_style(table_a)

    doc.add_paragraph("")

    # ── PART B - EMPLOYMENT DETAILS ──
    _add_paragraph(doc, "PART B - EMPLOYMENT DETAILS", bold=True, size=12, space_after=6)

    table_b = doc.add_table(rows=5, cols=2)
    table_b.style = "Table Grid"
    emp_type = incident.get("employment_type", case_data.get("employment_type", ""))
    fields_b = [
        ("Basis of Employment", _fmt(emp_type, "[REQUIRED] Permanent / Casual / Contractor")),
        ("How long at this job?", _fmt(incident.get("tenure"), "[REQUIRED]")),
        ("Average hours/days per week", _fmt(case_data.get("shift_structure"), "[REQUIRED]")),
        ("Shift", _fmt(incident.get("shift_type"), "[REQUIRED] Day / Afternoon / Night")),
        ("Time shift started", _fmt(incident.get("shift_start_time"), "[REQUIRED]")),
    ]
    for i, (label, val) in enumerate(fields_b):
        _set_cell_text(table_b.cell(i, 0), label, bold=True, size=10)
        _set_cell_text(table_b.cell(i, 1), val, size=10)
    _set_table_style(table_b)

    doc.add_paragraph("")

    # ── PART C - INJURY DETAILS ──
    _add_paragraph(doc, "PART C - INJURY DETAILS", bold=True, size=12, space_after=6)

    table_c = doc.add_table(rows=5, cols=2)
    table_c.style = "Table Grid"
    fields_c = [
        ("Nature of injury", _fmt(incident.get("nature_of_injury"), _fmt(case_data.get("injury_description")))),
        ("Body location of injury", _fmt(incident.get("body_part"), "[REQUIRED]")),
        ("Injury Treatment", _fmt(incident.get("treatment_level"), "[REQUIRED] No treatment / First Aid / Doctor / Hospital")),
        ("Is this a lost time injury?", _fmt(incident.get("lost_time"), "[REQUIRED] Y / N")),
        ("Is a Workers Compensation Claim being made?", _fmt(incident.get("claim_made"), "[REVIEW] Y")),
    ]
    for i, (label, val) in enumerate(fields_c):
        _set_cell_text(table_c.cell(i, 0), label, bold=True, size=10)
        _set_cell_text(table_c.cell(i, 1), val, size=10)
    _set_table_style(table_c)

    doc.add_paragraph("")

    # ── PART D - ACKNOWLEDGMENT ──
    _add_paragraph(doc, "PART D - ACKNOWLEDGMENT OF INJURY / DATE OF ENTRY", bold=True, size=12, space_after=6)

    table_d = doc.add_table(rows=3, cols=3)
    table_d.style = "Table Grid"
    _set_cell_text(table_d.cell(0, 0), "", size=10)
    _set_cell_text(table_d.cell(0, 1), "Signature", bold=True, size=10)
    _set_cell_text(table_d.cell(0, 2), "Date", bold=True, size=10)
    _set_cell_text(table_d.cell(1, 0), "Employee", bold=True, size=10)
    _set_cell_text(table_d.cell(1, 1), "", size=10)
    _set_cell_text(table_d.cell(1, 2), "", size=10)
    _set_cell_text(table_d.cell(2, 0), "Employer", bold=True, size=10)
    _set_cell_text(table_d.cell(2, 1), "", size=10)
    _set_cell_text(table_d.cell(2, 2), "", size=10)
    _set_table_style(table_d)

    doc.add_paragraph("")

    # ── PART E - INVESTIGATION ──
    _add_paragraph(doc, "PART E - ACCIDENT / INCIDENT INVESTIGATION", bold=True, size=12, space_after=6)

    p = _add_paragraph(doc, "", size=10, space_after=6)
    _marker_run(p, "[REVIEW] ")
    p.add_run("The following investigation section requires manual completion based on site inspection and interviews.").font.size = Pt(10)

    _add_paragraph(doc, "Contributing Factors to Consider:", bold=True, size=11, space_after=4)

    # Person factors
    _add_paragraph(doc, "Person Factors:", bold=True, size=10, space_after=4)
    table_ep = doc.add_table(rows=11, cols=4)
    table_ep.style = "Table Grid"
    _set_cell_text(table_ep.cell(0, 0), "Factor", bold=True, size=9)
    _set_cell_text(table_ep.cell(0, 1), "Y", bold=True, size=9)
    _set_cell_text(table_ep.cell(0, 2), "N", bold=True, size=9)
    _set_cell_text(table_ep.cell(0, 3), "N/A", bold=True, size=9)
    person_factors = [
        "Aware of the hazard", "Experienced at the task", "Familiar with the work area",
        "Inducted to the site / task", "Using appropriate PPE", "Was training provided",
        "Supervision provided", "Job Analysis performed", "Task not modified / changed",
        "PPE provided",
    ]
    for i, factor in enumerate(person_factors):
        _set_cell_text(table_ep.cell(i + 1, 0), factor, size=9)
        _set_cell_text(table_ep.cell(i + 1, 1), "", size=9)
        _set_cell_text(table_ep.cell(i + 1, 2), "", size=9)
        _set_cell_text(table_ep.cell(i + 1, 3), "", size=9)
    _set_table_style(table_ep)

    doc.add_paragraph("")
    _add_paragraph(doc, "Environment Factors:", bold=True, size=10, space_after=4)
    table_ee = doc.add_table(rows=7, cols=4)
    table_ee.style = "Table Grid"
    _set_cell_text(table_ee.cell(0, 0), "Factor", bold=True, size=9)
    _set_cell_text(table_ee.cell(0, 1), "Y", bold=True, size=9)
    _set_cell_text(table_ee.cell(0, 2), "N", bold=True, size=9)
    _set_cell_text(table_ee.cell(0, 3), "N/A", bold=True, size=9)
    env_factors = [
        "Adequate temperature conditions", "Adequate lighting", "Adequate working space",
        "Clear floor and walkways", "Adequate housekeeping", "Safe noise level",
    ]
    for i, factor in enumerate(env_factors):
        _set_cell_text(table_ee.cell(i + 1, 0), factor, size=9)
        _set_cell_text(table_ee.cell(i + 1, 1), "", size=9)
        _set_cell_text(table_ee.cell(i + 1, 2), "", size=9)
        _set_cell_text(table_ee.cell(i + 1, 3), "", size=9)
    _set_table_style(table_ee)

    doc.add_paragraph("")
    _add_paragraph(doc, "Equipment Factors:", bold=True, size=10, space_after=4)
    table_eq = doc.add_table(rows=7, cols=4)
    table_eq.style = "Table Grid"
    _set_cell_text(table_eq.cell(0, 0), "Factor", bold=True, size=9)
    _set_cell_text(table_eq.cell(0, 1), "Y", bold=True, size=9)
    _set_cell_text(table_eq.cell(0, 2), "N", bold=True, size=9)
    _set_cell_text(table_eq.cell(0, 3), "N/A", bold=True, size=9)
    equip_factors = [
        "Correct equipment used", "Equipment in correct location", "Equipment guarded",
        "Preventative maintenance complete", "Equipment working properly",
        "Equipment had not been modified",
    ]
    for i, factor in enumerate(equip_factors):
        _set_cell_text(table_eq.cell(i + 1, 0), factor, size=9)
        _set_cell_text(table_eq.cell(i + 1, 1), "", size=9)
        _set_cell_text(table_eq.cell(i + 1, 2), "", size=9)
        _set_cell_text(table_eq.cell(i + 1, 3), "", size=9)
    _set_table_style(table_eq)

    doc.add_paragraph("")

    # ── PART F - RECOMMENDATIONS ──
    _add_paragraph(doc, "PART F - RECOMMENDATIONS", bold=True, size=12, space_after=6)
    p = _add_paragraph(doc, "", size=10, space_after=6)
    _marker_run(p, "[REVIEW] ")
    p.add_run("Complete based on investigation findings.").font.size = Pt(10)

    table_f = doc.add_table(rows=6, cols=3)
    table_f.style = "Table Grid"
    _set_cell_text(table_f.cell(0, 0), "Control", bold=True, size=9)
    _set_cell_text(table_f.cell(0, 1), "Y/N", bold=True, size=9)
    _set_cell_text(table_f.cell(0, 2), "Why/How", bold=True, size=9)
    recs = [
        "Can the risk be eliminated?",
        "Can equipment or materials be substituted?",
        "Can engineering solutions be adopted?",
        "Can administrative controls be developed?",
        "Is PPE required?",
    ]
    for i, rec in enumerate(recs):
        _set_cell_text(table_f.cell(i + 1, 0), rec, size=9)
        _set_cell_text(table_f.cell(i + 1, 1), "", size=9)
        _set_cell_text(table_f.cell(i + 1, 2), "", size=9)
    _set_table_style(table_f)

    doc.add_paragraph("")

    # ── PART G - ACTION PLAN ──
    _add_paragraph(doc, "PART G - ACTION PLAN", bold=True, size=12, space_after=6)
    p = _add_paragraph(doc, "", size=10, space_after=6)
    _marker_run(p, "[REQUIRED] ")
    p.add_run("Complete action plan with recommended corrective actions.").font.size = Pt(10)

    table_g = doc.add_table(rows=4, cols=3)
    table_g.style = "Table Grid"
    _set_cell_text(table_g.cell(0, 0), "Recommended Actions", bold=True, size=9)
    _set_cell_text(table_g.cell(0, 1), "Implementation Date", bold=True, size=9)
    _set_cell_text(table_g.cell(0, 2), "Responsibility", bold=True, size=9)
    for i in range(1, 4):
        for j in range(3):
            _set_cell_text(table_g.cell(i, j), "", size=9)
    _set_table_style(table_g)

    doc.add_paragraph("")

    # ── PART H - COMPLETION ──
    _add_paragraph(doc, "PART H - COMPLETION / SIGNATURES", bold=True, size=12, space_after=6)
    table_h = doc.add_table(rows=3, cols=3)
    table_h.style = "Table Grid"
    _set_cell_text(table_h.cell(0, 0), "", size=10)
    _set_cell_text(table_h.cell(0, 1), "Signature", bold=True, size=10)
    _set_cell_text(table_h.cell(0, 2), "Date", bold=True, size=10)
    _set_cell_text(table_h.cell(1, 0), "Employee", bold=True, size=10)
    _set_cell_text(table_h.cell(2, 0), "Employer", bold=True, size=10)
    _set_table_style(table_h)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def generate_letter_to_doctor(case_data, doctor_data=None):
    """
    Generate Letter to Doctor regarding Return to Work.
    Returns BytesIO with .docx.
    """
    doc = Document()
    doctor = doctor_data or {}
    state = case_data.get("state", "VIC")
    agent = AGENTS.get(state, AGENTS["VIC"])

    for section in doc.sections:
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Header fields
    table = doc.add_table(rows=3, cols=2)
    _set_cell_text(table.cell(0, 0), "Claim Number:", bold=True)
    _set_cell_text(table.cell(0, 1), _fmt(case_data.get("claim_number")))
    _set_cell_text(table.cell(1, 0), "Worker Name:", bold=True)
    _set_cell_text(table.cell(1, 1), _fmt(case_data.get("worker_name")))
    _set_cell_text(table.cell(2, 0), "Employer Name:", bold=True)
    _set_cell_text(table.cell(2, 1), EMPLOYER["short"])
    _set_table_style(table)

    doc.add_paragraph("")

    # Title
    _add_paragraph(doc, "Letter to GP/Physio/Psychologist re RTW", bold=True, size=13,
                   alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    # Date
    _add_paragraph(doc, date.today().strftime("%d/%m/%Y"), size=11, space_after=12)

    # Doctor address
    doctor_name = _fmt(doctor.get("doctor_name"), "[REQUIRED - Doctor Name]")
    doctor_address = _fmt(doctor.get("doctor_address"), "[REQUIRED - Doctor Address]")
    _add_paragraph(doc, doctor_name, size=11, space_after=2)
    _add_paragraph(doc, doctor_address, size=11, space_after=12)

    _add_paragraph(doc, f"Dear {doctor_name},", size=11, space_after=12)

    # Body
    worker_name = _fmt(case_data.get("worker_name"))
    _add_paragraph(doc,
        f"Re: Supporting your patient's recovery and return to work - {worker_name}",
        bold=True, size=11, space_after=12)

    _add_paragraph(doc,
        f"{EMPLOYER['short']} is committed to supporting the recovery of {worker_name} and their "
        f"return to safe, suitable and sustainable employment. As per the signed authority on the claim form, "
        f"we have developed return to work arrangements for your patient and would appreciate your review.",
        size=11, space_after=12)

    _add_paragraph(doc, "What do I need to do?", bold=True, size=12, space_after=6)

    actions = [
        "Review the Return to Work Arrangements attached",
        "Discuss the Return to Work Arrangements with your patient",
        "Provide feedback on the Return to Work Arrangements",
        "Sign second page if satisfied and email/mail a copy back to us",
    ]
    for action in actions:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(action)
        run.font.size = Pt(11)

    doc.add_paragraph("")

    _add_paragraph(doc,
        "If you have any questions or concerns, please don't hesitate to contact me.",
        size=11, space_after=12)

    _add_paragraph(doc, "Kind regards,", size=11, space_after=6)
    doc.add_paragraph("")
    _add_paragraph(doc, RTW_COORDINATOR["name"], bold=True, size=11, space_after=2)
    _add_paragraph(doc, RTW_COORDINATOR["role"], size=10, space_after=2)
    _add_paragraph(doc, EMPLOYER["name"], size=10, space_after=2)
    _add_paragraph(doc, f"Phone: {RTW_COORDINATOR['phone']}", size=10, space_after=2)
    _add_paragraph(doc, f"Email: {RTW_COORDINATOR['email']}", size=10, space_after=2)
    _add_paragraph(doc, RTW_COORDINATOR["address"], size=10, space_after=12)

    _add_paragraph(doc, "encl. Signed authority on the Worker Injury Claim Form", size=9, space_after=2)

    claims_manager = _fmt(doctor.get("claims_manager"), f"[REVIEW] Claims Manager, {agent['name']}")
    _add_paragraph(doc, f"cc: {claims_manager}", size=9, space_after=2)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def generate_rtw_arrangement(case_data, medical_data=None):
    """
    Generate Return to Work Arrangement document.
    Returns BytesIO with .docx.
    """
    doc = Document()
    med = medical_data or {}

    for section in doc.sections:
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    _add_paragraph(doc, "Return to Work Arrangement", bold=True, size=16,
                   alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    # Worker details
    _add_paragraph(doc, "Worker Details", bold=True, size=13, space_after=6)
    table_w = doc.add_table(rows=6, cols=2)
    table_w.style = "Table Grid"
    worker_fields = [
        ("Worker Name", _fmt(case_data.get("worker_name"))),
        ("Claim Number", _fmt(case_data.get("claim_number"))),
        ("Date of Injury", _fmt_date(case_data.get("date_of_injury"))),
        ("Nature of Injury", _fmt(case_data.get("injury_description"))),
        ("Employer", EMPLOYER["short"]),
        ("Worksite", _fmt(case_data.get("site"))),
    ]
    for i, (label, val) in enumerate(worker_fields):
        _set_cell_text(table_w.cell(i, 0), label, bold=True, size=10)
        _set_cell_text(table_w.cell(i, 1), val, size=10)
    _set_table_style(table_w)

    doc.add_paragraph("")

    # Current capacity
    _add_paragraph(doc, "Current Capacity & Restrictions", bold=True, size=13, space_after=6)
    capacity = case_data.get("current_capacity", "Unknown")
    restrictions = _fmt(med.get("restrictions"), "[REVIEW] As per current Certificate of Capacity")

    table_c = doc.add_table(rows=4, cols=2)
    table_c.style = "Table Grid"
    cap_fields = [
        ("Current Capacity", _fmt(capacity)),
        ("Certified Hours", _fmt(med.get("certified_hours"), _fmt(case_data.get("shift_structure"), "[REQUIRED]"))),
        ("Certificate Period", f"{_fmt_date(med.get('cert_from'))} to {_fmt_date(med.get('cert_to'))}"),
        ("Medical Restrictions", restrictions),
    ]
    for i, (label, val) in enumerate(cap_fields):
        _set_cell_text(table_c.cell(i, 0), label, bold=True, size=10)
        _set_cell_text(table_c.cell(i, 1), val, size=10)
    _set_table_style(table_c)

    doc.add_paragraph("")

    # Suitable duties
    level = _get_suitable_level(capacity)
    duties_info = SUITABLE_DUTIES[level]

    _add_paragraph(doc, "Proposed Suitable Duties", bold=True, size=13, space_after=6)

    p = _add_paragraph(doc, "", size=10, space_after=6)
    _marker_run(p, "[REVIEW] ")
    p.add_run("The following duties are proposed based on current capacity. Please review and adjust as needed.").font.size = Pt(10)

    _add_paragraph(doc, duties_info["title"], bold=True, size=11, space_after=4)
    _add_paragraph(doc, f"Purpose: {duties_info['purpose']}", size=10, space_after=6)

    _add_paragraph(doc, "Duties may include:", bold=True, size=10, space_after=4)
    for duty in duties_info["duties"]:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(duty)
        run.font.size = Pt(10)

    doc.add_paragraph("")
    _add_paragraph(doc, "Restrictions:", bold=True, size=10, space_after=4)
    for restriction in duties_info["restrictions"]:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(restriction)
        run.font.size = Pt(10)

    if duties_info["rest_break"]:
        doc.add_paragraph("")
        _add_paragraph(doc, f"Rest breaks: {duties_info['rest_break']}", size=10, space_after=6)

    doc.add_paragraph("")

    # Signatures
    _add_paragraph(doc, "Agreement", bold=True, size=13, space_after=6)
    _add_paragraph(doc,
        "I have read and agree to the above Return to Work Arrangement. "
        "I understand that this arrangement may be modified based on medical advice and progress.",
        size=10, space_after=12)

    table_s = doc.add_table(rows=4, cols=3)
    table_s.style = "Table Grid"
    _set_cell_text(table_s.cell(0, 0), "", size=10)
    _set_cell_text(table_s.cell(0, 1), "Signature", bold=True, size=10)
    _set_cell_text(table_s.cell(0, 2), "Date", bold=True, size=10)
    _set_cell_text(table_s.cell(1, 0), "Worker", bold=True, size=10)
    _set_cell_text(table_s.cell(2, 0), "Employer", bold=True, size=10)
    _set_cell_text(table_s.cell(3, 0), "Treating Practitioner", bold=True, size=10)
    _set_table_style(table_s)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def generate_rtw_plan(case_data, medical_data=None):
    """
    Generate Return to Work Plan with 4-week progressive schedule.
    Returns BytesIO with .docx.
    """
    doc = Document()
    med = medical_data or {}
    state = case_data.get("state", "VIC")

    for section in doc.sections:
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    _add_paragraph(doc, "Return to Work Plan", bold=True, size=16,
                   alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    _add_paragraph(doc, f"Employer: {EMPLOYER['name']}", size=11,
                   alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    # ── EMPLOYER DETAILS ──
    _add_paragraph(doc, "EMPLOYER DETAILS", bold=True, size=12, space_after=6)
    table_emp = doc.add_table(rows=5, cols=2)
    table_emp.style = "Table Grid"
    emp_fields = [
        ("Company Name", EMPLOYER["name"]),
        ("Prepared by", f"{RTW_COORDINATOR['name']}, {RTW_COORDINATOR['role']}"),
        ("Contact Number", RTW_COORDINATOR["phone"]),
        ("Plan number", _fmt(med.get("plan_number"), "[REVIEW] 1")),
        ("To be reviewed", _fmt_date(med.get("review_date"),
            (date.today() + timedelta(days=28)).strftime("%d/%m/%Y"))),
    ]
    for i, (label, val) in enumerate(emp_fields):
        _set_cell_text(table_emp.cell(i, 0), label, bold=True, size=10)
        _set_cell_text(table_emp.cell(i, 1), val, size=10)
    _set_table_style(table_emp)

    doc.add_paragraph("")

    # ── WORKER DETAILS ──
    _add_paragraph(doc, "WORKER DETAILS", bold=True, size=12, space_after=6)
    table_wrk = doc.add_table(rows=8, cols=2)
    table_wrk.style = "Table Grid"
    wrk_fields = [
        ("Name", _fmt(case_data.get("worker_name"))),
        ("Place of residence", _fmt(med.get("worker_address"), "[REQUIRED]")),
        ("Telephone", _fmt(med.get("worker_phone"), "[REQUIRED]")),
        ("Date of birth", _fmt(med.get("worker_dob"), "[REQUIRED]")),
        ("Occupation / pre-injury duties", _fmt(med.get("occupation"), "[REVIEW] Cleaner")),
        ("Date of injury", _fmt_date(case_data.get("date_of_injury"))),
        ("Claim number", _fmt(case_data.get("claim_number"))),
        ("Interpreter required?", _fmt(med.get("interpreter_needed"), "[REQUIRED] Yes / No")),
    ]
    for i, (label, val) in enumerate(wrk_fields):
        _set_cell_text(table_wrk.cell(i, 0), label, bold=True, size=10)
        _set_cell_text(table_wrk.cell(i, 1), val, size=10)
    _set_table_style(table_wrk)

    doc.add_paragraph("")

    # ── TREATING PRACTITIONER ──
    _add_paragraph(doc, "TREATING MEDICAL PRACTITIONER", bold=True, size=12, space_after=6)
    table_doc = doc.add_table(rows=6, cols=2)
    table_doc.style = "Table Grid"
    doc_fields = [
        ("Name", _fmt(med.get("doctor_name"), "[REQUIRED]")),
        ("Doctor has been contacted", "[REVIEW] Y / N"),
        ("Address", _fmt(med.get("doctor_address"), "[REQUIRED]")),
        ("Telephone", _fmt(med.get("doctor_phone"), "[REQUIRED]")),
        ("Fax", _fmt(med.get("doctor_fax"), "[REQUIRED]")),
        ("Nature of injury", _fmt(case_data.get("injury_description"))),
    ]
    for i, (label, val) in enumerate(doc_fields):
        _set_cell_text(table_doc.cell(i, 0), label, bold=True, size=10)
        _set_cell_text(table_doc.cell(i, 1), val, size=10)
    _set_table_style(table_doc)

    doc.add_paragraph("")

    # ── CURRENT RTW RESTRICTIONS ──
    _add_paragraph(doc, "CURRENT RETURN TO WORK RESTRICTIONS", bold=True, size=12, space_after=6)

    # Calculate hours
    hours_per_day = med.get("hours_per_day")
    days_per_week = med.get("days_per_week")
    current_hours = None
    if hours_per_day and days_per_week:
        try:
            current_hours = float(hours_per_day) * int(days_per_week)
        except (ValueError, TypeError):
            pass

    pre_injury_hours = med.get("pre_injury_hours", 38)

    table_r = doc.add_table(rows=5, cols=2)
    table_r.style = "Table Grid"
    r_fields = [
        ("Current capacity", _fmt(case_data.get("current_capacity"))),
        ("Certified hours", _fmt(case_data.get("shift_structure"),
            f"[REVIEW] {current_hours} hrs/week" if current_hours else "[REQUIRED]")),
        ("Pre-injury average", _fmt(med.get("pre_injury_hours"), "[REVIEW] 38 hrs/week")),
        ("Medical Constraints", _fmt(med.get("restrictions"), "[REVIEW] As per Certificate of Capacity")),
        ("Current Certificate", f"{_fmt_date(med.get('cert_from'))} to {_fmt_date(med.get('cert_to'))}"),
    ]
    for i, (label, val) in enumerate(r_fields):
        _set_cell_text(table_r.cell(i, 0), label, bold=True, size=10)
        _set_cell_text(table_r.cell(i, 1), val, size=10)
    _set_table_style(table_r)

    doc.add_paragraph("")

    # ── RETURN TO WORK SECTION ──
    _add_paragraph(doc, "RETURN TO WORK", bold=True, size=12, space_after=6)

    level = _get_suitable_level(case_data.get("current_capacity"))
    duties_info = SUITABLE_DUTIES[level]

    table_rtw = doc.add_table(rows=5, cols=2)
    table_rtw.style = "Table Grid"
    rtw_fields = [
        ("Return to work position", _fmt(med.get("rtw_position"), f"[REVIEW] Modified Duties - {duties_info['title']}")),
        ("Return to work location", _fmt(case_data.get("site"), "[REQUIRED]")),
        ("Return to Work Goals",
         f"[REVIEW] Progressive return to pre-injury duties over 4 weeks. "
         f"Commence at {duties_info['title']} and progress as tolerated."),
        ("Specific duties/tasks", "[REVIEW] " + "; ".join(duties_info["duties"][:4])),
        ("Restrictions", "; ".join(duties_info["restrictions"][:3])),
    ]
    for i, (label, val) in enumerate(rtw_fields):
        _set_cell_text(table_rtw.cell(i, 0), label, bold=True, size=10)
        _set_cell_text(table_rtw.cell(i, 1), val, size=10)
    _set_table_style(table_rtw)

    doc.add_paragraph("")

    # ── HOURS OF WORK TABLE (4-week progressive) ──
    _add_paragraph(doc, "HOURS OF WORK", bold=True, size=12, space_after=6)

    p = _add_paragraph(doc, "", size=10, space_after=6)
    _marker_run(p, "[REVIEW] ")
    p.add_run("The following 4-week schedule is auto-generated based on current capacity. Adjust start/finish times and duties levels as needed.").font.size = Pt(10)

    if not current_hours:
        current_hours = 3 * (days_per_week or 3)
    try:
        pre_injury_hours = float(pre_injury_hours)
    except (ValueError, TypeError):
        pre_injury_hours = 38.0

    schedule = _build_progressive_hours(current_hours, pre_injury_hours)

    table_hrs = doc.add_table(rows=5, cols=3)
    table_hrs.style = "Table Grid"
    _set_cell_text(table_hrs.cell(0, 0), "Week", bold=True, size=10)
    _set_cell_text(table_hrs.cell(0, 1), "Total Hours/Week", bold=True, size=10)
    _set_cell_text(table_hrs.cell(0, 2), "Duties Level", bold=True, size=10)

    # Map weeks to progressive duty levels
    for w in range(4):
        week_hrs = schedule[w]
        if week_hrs <= current_hours:
            week_level = level
        elif week_hrs >= pre_injury_hours * 0.9:
            week_level = 4
        else:
            week_level = min(level + w, 4)

        _set_cell_text(table_hrs.cell(w + 1, 0), f"Week {w + 1}", bold=True, size=10)
        _set_cell_text(table_hrs.cell(w + 1, 1), f"{week_hrs} hrs", size=10)
        _set_cell_text(table_hrs.cell(w + 1, 2), f"L{week_level}", size=10)
    _set_table_style(table_hrs)

    doc.add_paragraph("")

    # ── SIGNATURES ──
    _add_paragraph(doc, "SIGNATURES", bold=True, size=12, space_after=6)

    table_sig = doc.add_table(rows=4, cols=3)
    table_sig.style = "Table Grid"
    _set_cell_text(table_sig.cell(0, 0), "", size=10)
    _set_cell_text(table_sig.cell(0, 1), "Signature", bold=True, size=10)
    _set_cell_text(table_sig.cell(0, 2), "Date", bold=True, size=10)
    _set_cell_text(table_sig.cell(1, 0), "Employer", bold=True, size=10)
    _set_cell_text(table_sig.cell(2, 0), "Worker", bold=True, size=10)
    _set_cell_text(table_sig.cell(3, 0), "Treating Practitioner", bold=True, size=10)
    _set_table_style(table_sig)

    doc.add_paragraph("")

    # Occupational rehab
    _add_paragraph(doc, "Will assistance for RTW or other occupational rehab services be required?",
                   bold=True, size=10, space_after=4)
    _add_paragraph(doc, "[REVIEW] Yes / No", size=10, space_after=2)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ── Master list of available document types ────────────────────

AVAILABLE_DOCUMENTS = {
    "letter_to_worker": {
        "name": "Letter to Worker",
        "description": "Letter advising worker about Return to Work Arrangements",
        "needs_review": "Minimal",
        "generator": generate_letter_to_worker,
        "icon": "📄",
    },
    "rtw_information": {
        "name": "RTW Information",
        "description": "Important Return to Work Information (state-specific)",
        "needs_review": "No",
        "generator": generate_rtw_information,
        "icon": "📋",
    },
    "register_of_injury": {
        "name": "Register of Injury & Investigation",
        "description": "Full register including investigation, recommendations and action plan",
        "needs_review": "Yes - Parts E, F, G need manual completion",
        "generator": generate_register_of_injury,
        "icon": "📝",
    },
    "letter_to_doctor": {
        "name": "Letter to Doctor",
        "description": "Letter to treating practitioner regarding RTW arrangements",
        "needs_review": "Minimal",
        "generator": generate_letter_to_doctor,
        "icon": "🏥",
    },
    "rtw_arrangement": {
        "name": "RTW Arrangement",
        "description": "Return to Work Arrangement with proposed suitable duties",
        "needs_review": "Yes - review proposed duties",
        "generator": generate_rtw_arrangement,
        "icon": "📎",
    },
    "rtw_plan": {
        "name": "RTW Plan",
        "description": "Return to Work Plan with 4-week progressive schedule",
        "needs_review": "Yes - review schedule and duties",
        "generator": generate_rtw_plan,
        "icon": "📊",
    },
}


def generate_documents(case_data, doc_types, medical_data=None, doctor_data=None, incident_data=None):
    """
    Generate multiple documents for a case.

    Args:
        case_data: dict with case fields (worker_name, claim_number, etc.)
        doc_types: list of document type keys from AVAILABLE_DOCUMENTS
        medical_data: dict with medical/COC details
        doctor_data: dict with doctor details
        incident_data: dict with incident details

    Returns:
        dict mapping doc_type key -> (filename, BytesIO)
    """
    results = {}
    worker_name = (case_data.get("worker_name") or "Worker").replace(" ", "_")
    today_str = date.today().strftime("%Y%m%d")

    for doc_type in doc_types:
        if doc_type not in AVAILABLE_DOCUMENTS:
            continue

        info = AVAILABLE_DOCUMENTS[doc_type]
        gen_func = info["generator"]

        if doc_type == "register_of_injury":
            buf = gen_func(case_data, incident_data)
        elif doc_type == "letter_to_doctor":
            buf = gen_func(case_data, doctor_data)
        elif doc_type in ("rtw_arrangement", "rtw_plan"):
            buf = gen_func(case_data, medical_data)
        else:
            buf = gen_func(case_data)

        filename = f"{worker_name}_{info['name'].replace(' ', '_')}_{today_str}.docx"
        results[doc_type] = (filename, buf)

    return results
